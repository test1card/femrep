"""ГОСТ 7.32-2017 Russian DOCX renderer: structure, formatting, and the
zero-English-label guarantee. No Ansys/Qt needed."""
from __future__ import annotations
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

# English label words that must NOT survive translation. Proper nouns
# (Ansys/Nastran/femis), the abbreviation GCI, file names and unit symbols are
# intentionally allowed and excluded from this list.
FORBIDDEN = ["Summary", "Model", "Meshing", "Composites", "Mechanical", "Results",
             "Governance", "Manifest", "Field", "Value", "Element", "Count", "Metric",
             "Criterion", "Conclusion", "Note", "Evidence", "Status", "Gate", "PASS",
             "FAIL", "Convergence", "Units", "Mesh", "Report", "Solver", "Platform",
             "Deck", "Readiness", "Range", "Customer", "Project", "Prepared", "Nodes",
             "Analysis", "Temperature", "stress", "displacement"]


def _fixture():
    results = {
        "result_file": "000.f06", "result_sha256": "abc",
        "primary_qoi": {"name": "temperature", "units": "K", "min": 300.0, "max": 305.0,
                        "hot_node": 1, "cold_node": 2},
        "qoi_catalog": [{"name": "temperature"}],
        "mesh": {"nodes": 10, "elements": 4, "element_types": {"tet": 4}},
        "convergence": {"converged": True, "substeps": 1, "note": "ок"},
    }
    manifest = {"analysis_type": "thermal", "units": "SI", "solver": "Nastran",
                "solver_version": "0.1", "platform": "darwin", "command_line": "femrep ...",
                "deck_path": None, "superseded_by": None}
    checks = {"claim": "...", "gci": None,
              "gates": [{"gate": "units", "verdict": "pass", "note": ""},
                        {"gate": "convergence", "verdict": "pass", "note": ""},
                        {"gate": "mesh_independence_GCI", "verdict": "not_done", "note": ""}],
              "readiness": {"status": "issue_with_limitations", "summary": "",
                            "items": [{"key": "result_hash", "status": "complete", "note": ""},
                                      {"key": "deck_hash", "status": "missing", "note": ""}]}}
    return results, manifest, checks


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_gost_structural_elements_present():
    from femrep import report_gost_docx
    results, manifest, checks = _fixture()
    doc = report_gost_docx.build_gost_doc(results, manifest, checks, {}, {}, {"generated": "2026-06-30"})
    text = _all_text(doc)
    for struct in ("ОТЧЁТ О НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ", "РЕФЕРАТ", "СОДЕРЖАНИЕ",
                   "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ"):
        assert struct in text, f"missing structural element {struct!r}"
    # numbered основная часть sections, Russian
    assert re.search(r"\b1 Сводка результатов\b", text)
    assert "9 Протокол вычислений" in text


def test_gost_has_no_english_label_words():
    from femrep import report_gost_docx
    results, manifest, checks = _fixture()
    doc = report_gost_docx.build_gost_doc(results, manifest, checks, {}, {}, {"generated": "2026-06-30"})
    text = _all_text(doc)
    leaked = [w for w in FORBIDDEN if re.search(rf"\b{re.escape(w)}\b", text, re.IGNORECASE)]
    assert leaked == [], f"English label words leaked into the GOST report: {leaked}"


def test_gost_formatting_font_spacing_margins():
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_LINE_SPACING
    from femrep import report_gost_docx
    results, manifest, checks = _fixture()
    doc = report_gost_docx.build_gost_doc(results, manifest, checks, {}, {}, {"generated": "2026-06-30"})
    normal = doc.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal.font.size == Pt(14)
    assert normal.paragraph_format.line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE
    sec = doc.sections[0]
    # python-docx stores margins in twips, so compare in rounded mm (not exact EMU)
    assert round(sec.left_margin.mm) == 30 and round(sec.right_margin.mm) == 15
    assert round(sec.top_margin.mm) == 20 and round(sec.bottom_margin.mm) == 20
    assert sec.different_first_page_header_footer is True


def test_gost_honors_section_selection_and_verdict_mapping():
    from femrep import report_gost_docx, templates
    results, manifest, checks = _fixture()
    tpl = templates.default_template("g")
    tpl["sections"] = [{"key": "summary", "enabled": True, "intro": ""},
                       {"key": "governance", "enabled": True, "intro": ""}]
    cfg = templates.to_config(tpl)
    doc = report_gost_docx.build_gost_doc(results, manifest, checks, cfg, {}, {"generated": "2026-06-30"})
    # numbered основная-часть headings only (paragraphs starting "<n> ")
    headings = [p.text for p in doc.paragraphs if re.match(r"^\d+\s+\S", p.text)]
    assert headings == ["1 Сводка результатов", "2 Контроль качества (femis)"]
    assert "соответствует" in _all_text(doc)  # pass verdict mapped to Russian
