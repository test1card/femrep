"""femrep.report_gost_docx — ГОСТ 7.32-2017 отчёт о НИР, fully Russian, DOCX.

A dedicated renderer (the English report_pdf/report_docx are untouched). Produces
the ГОСТ structural elements — титульный лист, реферат, содержание, введение,
основная часть (numbered Russian sections from the template), заключение — with
ГОСТ formatting: Times New Roman 14, полуторный интервал, поля 30/15/20/20 мм,
абзацный отступ 1.25 см, выравнивание по ширине, сквозная нумерация страниц по
центру нижнего поля (титульный лист в счёт входит, номер не печатается).

All labels and prose come from locale_ru, so no English label words appear.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

from . import templates as _templates
from . import locale_ru as L

FONT = "Times New Roman"
SIZE = 14
# Predicted table count per section (keeps the реферат volume line honest without
# a two-pass render). Figures are counted from the files actually present.
_TABLES_PER_SECTION = {"summary": 1, "model": 1, "meshing": 1, "composites": 1,
                       "solve": 1, "results": 0, "gci": 1, "governance": 2, "manifest": 1}
_FIG_KEYS = {"results": ["contour_views", "contour", "deformed_shape", "time_history"],
             "gci": ["gci_convergence"]}


# --- low-level docx helpers --------------------------------------------------

def _field(paragraph, instr: str, placeholder: str = ""):
    """Insert a Word field (PAGE / NUMPAGES / TOC) so Word computes it on open."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = placeholder
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, sep, text, end):
        run._r.append(el)
    return run


def _set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SIZE)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), FONT)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(0)


def _page_setup(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.left_margin, sec.right_margin = Mm(30), Mm(15)
    sec.top_margin, sec.bottom_margin = Mm(20), Mm(20)
    sec.different_first_page_header_footer = True   # титульный лист без номера
    footer_p = sec.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(footer_p, "PAGE", "2")


def _p(doc, text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, indent=True,
       upper=False, size=SIZE):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    if text:
        run = para.add_run(text.upper() if upper else text)
        run.bold = bold
        run.font.size = Pt(size)
    return para


def _struct_heading(doc, text, first=False):
    """Centered, uppercase structural-element heading, each on a new page."""
    if not first:
        doc.add_page_break()
    return _p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False, upper=True)


def _section_heading(doc, n, text):
    return _p(doc, f"{n} {text}", align=WD_ALIGN_PARAGRAPH.LEFT, bold=True, indent=True)


def _table(doc, n_table, caption, headers, rows):
    _p(doc, f"Таблица {n_table} — {caption}", align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


# --- титульный лист ----------------------------------------------------------

def _titul(doc, cfg, manifest, meta):
    b = cfg
    org = b.get("company") or "Организация-исполнитель"
    ministry = b.get("ministry") or ""
    if ministry:
        _p(doc, ministry, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    _p(doc, org, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    udc = b.get("udc") or "—"
    _p(doc, f"{L.TITLE['udc']} {udc}", align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    for _ in range(2):
        _p(doc, "", indent=False)
    # гриф утверждения
    _p(doc, L.TITLE["approve"], align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, indent=False)
    _p(doc, f"{b.get('head_org_title') or L.TITLE['head_org']}",
       align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    _p(doc, f"_______________ {b.get('approved_by') or ''}",
       align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    _p(doc, f"«___» __________ {b.get('year') or meta.get('generated', '')[:4]} г.",
       align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
    for _ in range(3):
        _p(doc, "", indent=False)
    _p(doc, L.TITLE["report_kind"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    title = b.get("title") or "Отчёт о конечно-элементном расчёте"
    _p(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    kind = (L.TITLE["report_type_interim"] if (b.get("report_type") == "interim")
            else L.TITLE["report_type_final"])
    _p(doc, kind, align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    for _ in range(4):
        _p(doc, "", indent=False)
    # руководитель НИР
    _p(doc, f"{b.get('head_work_title') or L.TITLE['head_work']}"
            f"        _______________   {b.get('prepared_by') or ''}",
       align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    for _ in range(3):
        _p(doc, "", indent=False)
    city = b.get("city") or ""
    year = b.get("year") or meta.get("generated", "")[:4]
    _p(doc, f"{city} {year}".strip(), align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)


# --- основная часть: per-section Russian builders ----------------------------

def _g_summary(doc, n, ctx):
    r, m, checks = ctx["results"], ctx["manifest"], ctx["checks"]
    q = r["primary_qoi"]
    name = L.qoi_ru(q.get("name", "")); units = L.units_ru(q.get("units", ""))
    readiness = checks.get("readiness") or {}
    _table(doc, ctx["nt"](), "Сводные показатели расчёта",
           [L.HDR["param"], L.HDR["value"]], [
        ["Готовность", L.readiness_status_ru(readiness.get("status", ""), "—")],
        ["Контролируемая величина", f"{name} ({units})"],
        ["Диапазон", f"{q.get('min','')} … {q.get('max','')} {units}"],
        ["Заказчик", cfg_get(ctx, "customer")],
        ["Подготовил / проверил / утвердил",
         f"{cfg_get(ctx,'prepared_by')} / {cfg_get(ctx,'checked_by')} / {cfg_get(ctx,'approved_by')}"],
    ])


def _g_model(doc, n, ctx):
    r, m = ctx["results"], ctx["manifest"]
    et = ", ".join(f"{k}: {v}" for k, v in r["mesh"].get("element_types", {}).items())
    n_elem = sum(r["mesh"].get("element_types", {}).values()) or r["mesh"].get("elements", 0)
    _table(doc, ctx["nt"](), "Параметры расчётной модели",
           [L.HDR["param"], L.HDR["value"]], [
        ["Узлы / элементы", f"{r['mesh'].get('nodes',0)} / {n_elem}"],
        ["Типы элементов", et or "—"],
        ["Тип анализа", L.analysis_ru(m.get("analysis_type", ""))],
        ["Единицы измерения", L.units_ru(m.get("units", "")) or "СИ"],
        ["Файл результатов", Path(r.get("result_file", "")).name],
    ])


def _g_meshing(doc, n, ctx):
    r = ctx["results"]
    rows = [[k, v] for k, v in r["mesh"].get("element_types", {}).items()] or [["—", "—"]]
    _table(doc, ctx["nt"](), "Состав конечно-элементной сетки",
           [L.HDR["element"], L.HDR["count"]], rows)


def _g_composites(doc, n, ctx):
    _p(doc, "Раздел приводит контрольный пример слоистого композита для демонстрации "
            "методики оценки (матрица жёсткости, симметрия укладки, критерий первого "
            "разрушения слоя). Реальные данные подключаются при их наличии.")
    _table(doc, ctx["nt"](), "Контрольный пример композита",
           [L.HDR["param"], L.HDR["value"]],
           [["Укладка", "[0/90/0]"], ["Симметрия", "B ≈ 0"]])


def _g_solve(doc, n, ctx):
    r, m = ctx["results"], ctx["manifest"]
    c = r.get("convergence", {})
    verdict = {True: "решение сошлось", False: "маркёр остановки / несходимости",
               None: "журнал отсутствует / отложено"}.get(c.get("converged"), "—")
    _table(doc, ctx["nt"](), "Решение и сходимость",
           [L.HDR["param"], L.HDR["value"]], [
        ["Вывод по сходимости", verdict],
        ["Подшаги / выводы", str(c.get("substeps", "—"))],
        ["Примечание", c.get("note", "—") if _is_ru(c.get("note", "")) else "—"],
        ["Решатель", f"{m.get('solver','')} {m.get('solver_version','')}"],
    ])


def _g_results(doc, n, ctx):
    figures = ctx["figures"]
    _p(doc, "Поля контролируемой величины и история нагружения приведены на рисунках "
            "ниже (при наличии геометрии).")
    for key in _FIG_KEYS["results"]:
        fp = figures.get(key)
        if fp and Path(fp).exists():
            doc.add_picture(str(fp), width=Cm(15))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(f"Рисунок {ctx['nf']()} — поле результатов").italic = True


def _g_gci(doc, n, ctx):
    gci = (ctx["checks"].get("gci"))
    if gci:
        _table(doc, ctx["nt"](), "Оценка сеточной сходимости (GCI)",
               [L.HDR["metric"], L.HDR["value"]], [
            ["Отношение измельчения r21", f"{gci.get('r21',0):.3f}"],
            ["Наблюдаемый порядок p", f"{gci.get('observed_order_p',0):.3f}"],
            ["GCI (мелкая), %", f"{gci.get('gci_fine_pct',0):.3f}"],
            ["Вывод", L.verdict_ru(gci.get('verdict_level', ''))],
        ])
    else:
        _p(doc, "Исследование сеточной сходимости не предоставлено — результат "
                "получен на одной сетке (см. критерии качества).")


def _g_governance(doc, n, ctx):
    checks = ctx["checks"]
    readiness = checks.get("readiness") or {}
    _p(doc, "Критерии качества — каждый вывод вычислен, а не назначен для "
            "«прохождения».")
    _table(doc, ctx["nt"](), "Готовность свидетельств",
           [L.HDR["evidence"], L.HDR["state"], L.HDR["note"]],
           [[L.READINESS_NAMES.get(i["key"], i["key"]), L.status_ru(i["status"]),
             L.readiness_note_ru(i)] for i in readiness.get("items", [])] or [["—", "—", "—"]])
    _table(doc, ctx["nt"](), "Критерии качества (femis)",
           [L.HDR["criterion"], L.HDR["conclusion"], L.HDR["note"]],
           [[L.GATE_NAMES.get(g["gate"], g["gate"]), L.verdict_ru(g["verdict"]),
             L.gate_note_ru(g, checks.get("gci"))] for g in checks.get("gates", [])])


def _g_manifest(doc, n, ctx):
    r, m = ctx["results"], ctx["manifest"]
    _p(doc, f"Контрольная сумма файла результатов (SHA-256): {r.get('result_sha256','—')}",
       indent=False)
    _table(doc, ctx["nt"](), "Протокол вычислений",
           [L.HDR["param"], L.HDR["value"]], [
        ["Решатель / версия", f"{m.get('solver','')} {m.get('solver_version','')}"],
        ["Расчётная модель", Path(m.get("deck_path")).name if m.get("deck_path") else "не предоставлена"],
        ["Дата формирования", (ctx["meta"].get("generated", "")[:10] if ctx.get("meta") else "")],
    ])


_GOST_BUILDERS = {
    "summary": _g_summary, "model": _g_model, "meshing": _g_meshing,
    "composites": _g_composites, "solve": _g_solve, "results": _g_results,
    "gci": _g_gci, "governance": _g_governance, "manifest": _g_manifest,
}


def cfg_get(ctx, key):
    return ctx["cfg"].get(key) or "—"


def _is_ru(text: str) -> bool:
    import re
    return bool(re.search(r"[А-Яа-я]", text or ""))


# --- assembly ----------------------------------------------------------------

def _enabled_sections(cfg):
    if "sections" in cfg:
        return [s for s in cfg["sections"]]
    return [{"key": k, "title": L.SECTION_TITLES_RU[k], "intro": ""} for k, _ in _templates.SECTIONS]


def build_gost_doc(results: dict, manifest: dict, checks: dict, cfg: dict,
                   figures: dict, meta: dict):
    doc = Document()
    _set_base_style(doc)
    _page_setup(doc)

    sections = _enabled_sections(cfg)
    # реферат volume counts (figures from present files; tables predicted)
    n_tables = sum(_TABLES_PER_SECTION.get(s["key"], 0) for s in sections if s.get("enabled", True))
    n_figs = sum(1 for s in sections if s.get("enabled", True)
                 for k in _FIG_KEYS.get(s["key"], []) if (figures or {}).get(k)
                 and Path(figures[k]).exists())

    _titul(doc, cfg, manifest, meta)

    # РЕФЕРАТ
    _struct_heading(doc, L.STRUCT["referat"])
    vol = _p(doc, "Отчёт ", indent=False)
    _field(vol, "NUMPAGES", "—")
    vol.add_run(f" с., {n_figs} рис., {n_tables} табл., 0 источн.")
    _p(doc, "Ключевые слова: " + ", ".join(L.keywords(results, manifest)), indent=False)
    _p(doc, L.referat_text(results, manifest))

    # СОДЕРЖАНИЕ (Word fills the field on open / update)
    _struct_heading(doc, L.STRUCT["soderzhanie"])
    toc = _p(doc, "", indent=False)
    _field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Обновите поле (F9), чтобы построить содержание.")

    # ВВЕДЕНИЕ
    _struct_heading(doc, L.STRUCT["vvedenie"])
    _p(doc, L.vvedenie_text(results, manifest))

    # ОСНОВНАЯ ЧАСТЬ
    counters = {"t": 0, "f": 0}
    def next_t():
        counters["t"] += 1; return counters["t"]
    def next_f():
        counters["f"] += 1; return counters["f"]
    ctx = {"results": results, "manifest": manifest, "checks": checks, "cfg": cfg,
           "figures": figures or {}, "meta": meta or {}, "nt": next_t, "nf": next_f}
    n = 0
    for s in sections:
        if not s.get("enabled", True):
            continue
        builder = _GOST_BUILDERS.get(s["key"])
        if builder is None:
            continue
        n += 1
        if n == 1:
            doc.add_page_break()
        _section_heading(doc, n, L.SECTION_TITLES_RU[s["key"]])
        if s.get("intro"):
            _p(doc, s["intro"])
        builder(doc, n, ctx)
    if n == 0:
        _p(doc, "Разделы основной части не выбраны в шаблоне.")

    # ЗАКЛЮЧЕНИЕ
    _struct_heading(doc, L.STRUCT["zaklyuchenie"])
    _p(doc, L.zaklyuchenie_text(results, checks))
    return doc


def render(results: dict, manifest: dict, checks: dict, cfg: dict,
           figures: dict, meta: dict, out_path: Path) -> None:
    build_gost_doc(results, manifest, checks, cfg, figures, meta).save(str(out_path))
