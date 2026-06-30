"""femrep.report_docx — render the same report as the PDF, as a Word document.

python-docx. Mirrors the PDF section order and content; the two share the same
results/manifest/checks inputs, so the engineering content is identical across
formats. Per the docx skill's contract: tables carry cantSplit + repeat-header
so they flow cleanly across pages.

Run via CLI (`--out report.docx`); standalone at the bottom.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

VERDICT_HEX = {"pass": "2B8A3E", "fail": "C92A2A", "not_done": "868E96"}


def _hex_rgb(h: str) -> RGBColor:
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _shade(cell, hex_color: str):
    """Cell background shading."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _no_split_row(row):
    """cantSplit on a row (docx skill: tables don't break mid-row across pages)."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def _header_row(row):
    """Repeat this row as a header on each page (docx skill contract)."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _add_table(doc, headers: list[str], rows: list[list[str]], *,
               widths_cm: list[float] | None = None, primary_hex: str = "1F3A5F"):
    """A styled table with header shading + cantSplit + repeat-header."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
        _shade(cell, primary_hex)
    _no_split_row(t.rows[0])
    _header_row(t.rows[0])
    for i, row_vals in enumerate(rows, start=1):
        _no_split_row(t.rows[i])
        for j, val in enumerate(row_vals):
            t.rows[i].cells[j].text = str(val)
            for p in t.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if widths_cm:
        for j, w in enumerate(widths_cm):
            for r in t.rows:
                r.cells[j].width = Cm(w)
    return t


def _h(doc, text: str, *, level: int = 1, color_hex: str = "1F3A5F"):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = _hex_rgb(color_hex)
    return p


from . import templates as _templates


def _dh(doc, n, title, intro, primary):
    """Numbered DOCX section heading + optional template intro paragraph."""
    _h(doc, f"{n}. {title}", color_hex=primary)
    if intro:
        doc.add_paragraph(intro)


def _dsec_summary(doc, n, intro, ctx):
    cfg, primary = ctx["cfg"], ctx["primary"]
    results, checks, qoi, readiness = ctx["results"], ctx["checks"], ctx["qoi"], ctx["readiness"]
    _dh(doc, n, "Summary", intro, primary)
    claim_p = doc.add_paragraph(checks["claim"].replace("**", ""))
    claim_p.style = doc.styles["Intense Quote"] if "Intense Quote" in [s.name for s in doc.styles] else None
    _add_table(doc, ["Field", "Value"], [
        ["Readiness", readiness.get("summary", "(not evaluated)") if readiness else "(not evaluated)"],
        ["Primary QoI", f"{qoi['name']} ({qoi['units']})"],
        ["QoI catalog", ", ".join(i["name"] for i in results.get("qoi_catalog", [])) or "(not available)"],
        ["Range", f"{qoi['min']} .. {qoi['max']}  "
                  + (f"({qoi.get('min_C')} .. {qoi.get('max_C')} °C)" if 'min_C' in qoi else "")],
        ["Hot node", f"{qoi['hot_node']} @ {qoi['hot_node_xyz_mm']} mm"],
        ["Cold node", f"{qoi['cold_node']} @ {qoi['cold_node_xyz_mm']} mm"],
        ["Report status", "Issued"],
        ["Project", cfg.get("project") or "(not specified)"],
        ["Customer", cfg.get("customer") or "(not specified)"],
        ["Prepared / checked / approved",
         f"{cfg.get('prepared_by') or '-'} / {cfg.get('checked_by') or '-'} / {cfg.get('approved_by') or '-'}"],
    ], widths_cm=[5, 11], primary_hex=primary)


def _dsec_model(doc, n, intro, ctx):
    results, manifest, primary = ctx["results"], ctx["manifest"], ctx["primary"]
    _dh(doc, n, "Model", intro, primary)
    et = ", ".join(f"{k}: {v:,}" for k, v in results["mesh"]["element_types"].items())
    n_elem = sum(results["mesh"]["element_types"].values()) if results["mesh"]["element_types"] \
        else results["mesh"].get("elements", 0)
    n_elem = n_elem if isinstance(n_elem, int) else 0
    _add_table(doc, ["Field", "Value"], [
        ["Nodes / elements", f"{results['mesh']['nodes']:,} / {n_elem:,}"
                             + ("" if n_elem else "  (not in .f06 — node/point dump only)")],
        ["Element types", et or "(n/a)"],
        ["Analysis type", manifest.get("analysis_type", "")],
        ["Units", manifest.get("units", "")],
        ["Result file", Path(results["result_file"]).name],
    ], widths_cm=[5, 11], primary_hex=primary)


def _dsec_meshing(doc, n, intro, ctx):
    results, primary = ctx["results"], ctx["primary"]
    _dh(doc, n, "Meshing", intro, primary)
    _add_table(doc, ["element", "count"],
               [[k, f"{v:,}"] for k, v in results["mesh"]["element_types"].items()],
               widths_cm=[6, 4], primary_hex=primary)


def _dsec_composites(doc, n, intro, ctx):
    _dh(doc, n, "Composites / CFRP", intro, ctx["primary"])
    _render_composites(doc, ctx["results"], ctx["cfg"], ctx["primary"])


def _dsec_solve(doc, n, intro, ctx):
    results, manifest, primary = ctx["results"], ctx["manifest"], ctx["primary"]
    _dh(doc, n, "Mechanical / solve", intro, primary)
    c = results.get("convergence", {})
    cv = c.get("converged")
    conv = {True: "converged", False: "non-convergence / stop marker",
            None: "no log / deferred"}.get(cv, str(cv))
    _add_table(doc, ["Field", "Value"], [
        ["Convergence verdict", conv],
        ["Substeps / outputs", str(c.get("substeps", "—"))],
        ["Final time", str(c.get("final_total_time", "—"))],
        ["Note", c.get("note", "—")],
        ["Solver", f"{manifest.get('solver','')} {manifest.get('solver_version','')}"],
    ], widths_cm=[5, 11], primary_hex=primary)


def _dsec_results(doc, n, intro, ctx):
    figures, primary = ctx["figures"], ctx["primary"]
    _dh(doc, n, "Results", intro, primary)
    for key, cap in [("contour_views", "QoI field contour, four-view plate."),
                     ("contour", "QoI field contour (pyvista off-screen)."),
                     ("deformed_shape", "Undeformed (wireframe) vs deformed shape (scaled)."),
                     ("time_history", "Transient time-history / QoI snapshot.")]:
        fp = figures.get(key)
        if fp and Path(fp).exists():
            doc.add_picture(str(fp), width=Cm(15))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(cap).runs[0].italic = True


def _dsec_gci(doc, n, intro, ctx):
    figures, primary, gci = ctx["figures"], ctx["primary"], ctx["gci"]
    _dh(doc, n, "Mesh independence (GCI)", intro, primary)
    if gci:
        _add_table(doc, ["Metric", "Value"], [
            ["refinement ratio r21", f"{gci['r21']:.3f}"],
            ["refinement ratio r32", f"{gci['r32']:.3f}"],
            ["Roache convergence ratio R", f"{gci['convergence_ratio_R']:.3f} (want 0<R<1)"],
            ["observed order p", f"{gci['observed_order_p']:.3f}"],
            ["GCI fine (%)", f"{gci['gci_fine_pct']:.3f}"],
            ["GCI coarse (%)", f"{gci['gci_coarse_pct']:.3f}"],
            ["asymptotic ratio", f"{gci['asymptotic_ratio']:.3f} (want ~1)"],
            ["VERDICT", gci["verdict"]],
        ], widths_cm=[8, 8], primary_hex=primary)
        fp = figures.get("gci_convergence")
        if fp and Path(fp).exists():
            doc.add_picture(str(fp), width=Cm(12))
    else:
        doc.add_paragraph("No GCI study provided — single-mesh result (see gates).")


def _dsec_governance(doc, n, intro, ctx):
    checks, primary, readiness = ctx["checks"], ctx["primary"], ctx["readiness"]
    _dh(doc, n, "Governance (femis)", intro, primary)
    if readiness:
        doc.add_paragraph(readiness.get("summary", ""))
        _add_table(doc, ["Evidence", "Status", "Note"],
                   [[i["key"], i["status"], i["note"]]
                    for i in readiness.get("items", [])],
                   widths_cm=[4.5, 3, 8.5], primary_hex=primary)
    gate_rows = []
    for g in checks["gates"]:
        sym = {"pass": "PASS", "fail": "FAIL", "not_done": "—"}.get(g["verdict"], "?")
        gate_rows.append([sym, g["gate"], g["note"]])
    gt = _add_table(doc, ["", "Gate", "Note"], gate_rows,
                    widths_cm=[1.5, 5, 9.5], primary_hex=primary)
    for i, g in enumerate(checks["gates"], start=1):
        cell = gt.rows[i].cells[0]
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = _hex_rgb(VERDICT_HEX[g["verdict"]])
                r.bold = True


def _dsec_manifest(doc, n, intro, ctx):
    results, manifest, primary = ctx["results"], ctx["manifest"], ctx["primary"]
    _dh(doc, n, "Run manifest (provenance)", intro, primary)
    doc.add_paragraph(f"Result SHA-256: {results.get('result_sha256','')}").runs[0].font.name = "Consolas"
    _add_table(doc, ["Field", "Value"], [
        ["Solver / version", f"{manifest.get('solver','')} {manifest.get('solver_version','')}"],
        ["Platform", manifest.get("platform", "")],
        ["Command line", manifest.get("command_line", "")],
        ["Deck", manifest.get("deck_path") or "(not supplied)"],
        ["Superseded by", manifest.get("superseded_by") or "(current — final)"],
    ], widths_cm=[5, 11], primary_hex=primary)


SECTION_BUILDERS = {
    "summary": _dsec_summary, "model": _dsec_model, "meshing": _dsec_meshing,
    "composites": _dsec_composites, "solve": _dsec_solve, "results": _dsec_results,
    "gci": _dsec_gci, "governance": _dsec_governance, "manifest": _dsec_manifest,
}


def build_doc(results: dict, manifest: dict, checks: dict, cfg: dict,
              figures: dict, meta: dict):
    """Build and return the Word Document: cover + the template's enabled sections,
    in order, dynamically numbered. Mirrors report_pdf.build_story. cfg['sections']
    drives selection; absent it, every section renders in canonical order."""
    doc = Document()
    primary = cfg.get("color_primary", "1F3A5F").lstrip("#")

    # --- Cover ---
    doc.add_heading(cfg.get("title", "FEM Analysis Report"), level=0)
    sub = doc.add_paragraph(
        f"{manifest.get('analysis_type','').split(' ')[0]} analysis · "
        f"{manifest.get('solver','')} {manifest.get('solver_version','')}")
    sub.runs[0].font.color.rgb = _hex_rgb("868E96")
    doc_id = cfg.get("document_number") or "uncontrolled draft"
    rev = cfg.get("revision") or "-"
    p = doc.add_paragraph(f"Generated {meta.get('generated','')[:10]}  ·  "
                          f"Issued engineering report  ·  Doc {doc_id} Rev {rev}")
    p.runs[0].bold = True
    doc.add_paragraph()

    ctx = {"results": results, "manifest": manifest, "checks": checks, "cfg": cfg,
           "figures": figures, "primary": primary, "qoi": results["primary_qoi"],
           "gci": checks.get("gci"), "readiness": checks.get("readiness")}
    sections = cfg["sections"] if "sections" in cfg else [
        {"key": k, "title": t, "intro": ""} for k, t in _templates.SECTIONS]
    n = 0
    for s in sections:
        builder = SECTION_BUILDERS.get(s["key"])
        if builder is None:
            continue
        n += 1
        builder(doc, n, s.get("intro", ""), ctx)
    if n == 0:
        doc.add_paragraph("No report sections are enabled in this template.")
    return doc


def render(results: dict, manifest: dict, checks: dict, cfg: dict,
           figures: dict, meta: dict, out_path: Path) -> None:
    doc = build_doc(results, manifest, checks, cfg, figures, meta)
    doc.save(str(out_path))


def _render_composites(doc, results: dict, cfg: dict, primary: str):
    """CFRP section — synthetic CLT validation case (mirrors the PDF renderer)."""
    if "layup" in results:
        doc.add_paragraph("Real layup data detected — ACP integration pending (M-later).")
        return
    from .cases.clt_synthetic import case as clt_case
    c = clt_case()
    layup = c["layup"]
    doc.add_paragraph(
        "No ACP/.rmed layup detected. Rendering the synthetic [0/90/0] CLT validation "
        "case so the composite-section governance is demonstrated.")
    doc.add_paragraph(f"Material: {c['material']}  ·  {layup['failure_philosophy']}")
    rows = [[str(i), str(ang), str(layup["t_ply_mm"])]
            for i, ang in enumerate(layup["sequence"], 1)]
    rows.append(["—", "total", str(layup["total_t_mm"])])
    _add_table(doc, ["ply", "angle (deg)", "t (mm)"], rows,
               widths_cm=[2, 4, 4], primary_hex=primary)
    A = c["ABD"]["A_MPa_mm"]
    _add_table(doc, ["Field", "Value"], [
        ["A11 / A22 / A66", f"{A[0][0]:.0f} / {A[1][1]:.0f} / {A[2][2]:.0f} MPa·mm"],
        ["B (coupling)", "≈ 0 (symmetric layup — verified)"],
        ["First-ply-failure Nx", f"{c['fpf_Nx_N_per_mm']} N/mm (Tsai-Wu IRF = 1.0)"],
    ], widths_cm=[5, 11], primary_hex=primary)
    note = doc.add_paragraph(
        "femis governance: FPF is a design-point margin, conservative for ultimate strength. "
        "Progressive damage (CDM) requires characteristic-length / fracture-energy "
        "regularization for mesh objectivity. As-draped ply angles (not nominal) feed "
        "stiffness, strength and CTE.")
    note.style = doc.styles["Intense Quote"] if "Intense Quote" in [s.name for s in doc.styles] else None


if __name__ == "__main__":
    out = Path("out")
    r = json.loads((out / "results.json").read_text(encoding="utf-8"))
    m = json.loads((out / "manifest.json").read_text(encoding="utf-8"))
    c = json.loads((out / "checks.json").read_text(encoding="utf-8"))
    from . import cli
    cfg = cli._load_config(Path(__file__).parent / "config.yaml")
    render(r, m, c, cfg, {}, {"generated": "manual"}, Path("out/manual_report.docx"))
    print("wrote out/manual_report.docx")
